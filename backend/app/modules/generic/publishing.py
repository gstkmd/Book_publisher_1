import io
from app.modules.generic.models import Content

class PublishingService:
    async def export_docx(self, content: Content) -> bytes:
        try:
            from docx import Document
            from docx.shared import Inches
        except ImportError:
            print("python-docx not found. Please install dependencies.")
            return b"DOCX Mock Data"

        doc = Document()
        doc.add_heading(content.title, 0)
        
        import re
        import httpx
        from io import BytesIO

        # Simple HTML to DOCX converter for POC
        # In a real app, use a proper library like htmldocx
        body_html = str(content.body.get("text", "")) if isinstance(content.body, dict) else str(content.body)
        
        # Split by tags to handle images vs text
        parts = re.split(r'(<img[^>]+>)', body_html)
        
        async with httpx.AsyncClient() as client:
            for part in parts:
                if part.startswith('<img'):
                    # Extract src
                    src_match = re.search(r'src="([^"]+)"', part)
                    if src_match:
                        img_url = src_match.group(1)
                        try:
                            img_res = await client.get(img_url)
                            if img_res.status_code == 200:
                                img_data = BytesIO(img_res.content)
                                doc.add_picture(img_data, width=Inches(4))
                        except Exception as e:
                            print(f"Failed to download image {img_url}: {e}")
                else:
                    # Clean up other tags and add as paragraph
                    text_cleaned = re.sub('<[^<]+?>', '', part).strip()
                    if text_cleaned:
                        doc.add_paragraph(text_cleaned)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    async def export_epub(self, content: Content) -> bytes:
        try:
            from ebooklib import epub
        except ImportError:
            print("EbookLib not found.")
            return b"EPUB Mock Data"
        
        book = epub.EpubBook()
        book.set_identifier(str(content.id))
        book.set_title(content.title)
        book.set_language('en')
        book.add_author(f"User {content.author.ref.id}")
        
        # Create chapter
        c1 = epub.EpubHtml(title='Content', file_name='content.xhtml', lang='en')
        c1.content = f'<h1>{content.title}</h1><p>{str(content.body)}</p>'
        book.add_item(c1)
        
        # Define Table of Contents
        book.toc = (epub.Link('content.xhtml', 'Content', 'intro'), )
        
        # Add default NCX and Nav file
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        
        # Define CSS style
        style = 'body { font-family: Times, serif; }'
        nav_css = epub.EpubItem(uid="style_nav", file_name="style/nav.css", media_type="text/css", content=style)
        book.add_item(nav_css)
        
        # Basic spine
        book.spine = ['nav', c1]
        
        # Write to buffer
        # EbookLib writes to a file path or file-like object.
        # It's a bit tricky with BytesIO sometimes, but let's try.
        buffer = io.BytesIO()
        epub.write_epub(buffer, book, {})
        buffer.seek(0)
        return buffer.read()

publishing_service = PublishingService()
